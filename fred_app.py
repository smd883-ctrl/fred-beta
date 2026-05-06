"""
FRED — Families' Rights and Entitlements Directory
Beta Version 0.5

Complete build incorporating all agreed changes:
- Report language throughout (not audit)
- Lawful/unlawful throughout (not legal/illegal)
- Full landing page as entry point
- Get my report as single primary CTA
- Traffic light before upload
- Single upload zone with expander
- Sneak peek with email capture for beta
- Three tier engine: red #C0392B, amber #D4A017, green #1E8449
- Named accountable person amber only
- APDR connection in delivery log
- Lack of evidence is evidence of lack
- Correspondence module with transcript cross reference
- Post meeting summary generation
- Word and PDF document support
- Password protected document guidance
- School policy cross reference
- Annual review date capture
- Subscription signal specific to findings
- Full survey with notification opt-in
"""

import streamlit as st
import fitz
import re
import io
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.colors import HexColor
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.units import mm

# ─────────────────────────────────────────────
# CONSTANTS
# ─────────────────────────────────────────────

BRAND_BLUE = "#1B4F72"
BRAND_MID = "#2E86C1"
RED = "#C0392B"
AMBER = "#D4A017"
GREEN = "#1E8449"
RED_BG = "#FDEDEC"
AMBER_BG = "#FEF9E7"
GREEN_BG = "#EAFAF1"
BLUE_BG = "#EAF2FF"
GREY = "#717D7E"
PURPLE = "#8E44AD"
PURPLE_BG = "#F5EEF8"

# ─────────────────────────────────────────────
# PAGE CONFIG
# ─────────────────────────────────────────────

st.set_page_config(
    page_title="FRED — Families' Rights and Entitlements Directory",
    page_icon="📋",
    layout="centered",
    initial_sidebar_state="collapsed"
)

# ─────────────────────────────────────────────
# GLOBAL STYLES
# ─────────────────────────────────────────────

st.markdown(f"""
<style>
*{{box-sizing:border-box;}}
.main{{max-width:780px;margin:0 auto;}}

.fred-nav{{display:flex;justify-content:space-between;align-items:center;
    padding:14px 0;border-bottom:0.5px solid #D5D8DC;margin-bottom:0;}}
.fred-nav-logo{{font-size:15px;font-weight:500;letter-spacing:1px;color:{GREY};}}
.fred-nav-link{{font-size:13px;color:{GREY};cursor:pointer;}}
.fred-nav-links{{display:flex;gap:24px;}}

.fred-beta{{background:#FEF9E7;border-top:0.5px solid #F9CA24;
    border-bottom:0.5px solid #F9CA24;padding:9px 0;text-align:center;
    font-size:12px;color:#7D6608;}}

.fred-hero{{padding:52px 0 44px;text-align:center;}}
.fred-big{{font-size:80px;font-weight:500;letter-spacing:8px;
    color:{BRAND_BLUE};line-height:1;margin-bottom:6px;}}
.fred-full{{font-size:12px;letter-spacing:1.5px;color:{GREY};
    margin-bottom:24px;text-transform:uppercase;}}
.fred-hero-title{{font-size:26px;font-weight:500;color:#1A252F;
    line-height:1.25;margin-bottom:12px;}}
.fred-hero-title span{{color:{BRAND_BLUE};}}
.fred-hero-origin{{font-size:13px;color:{GREY};font-style:italic;margin-bottom:16px;}}
.fred-hero-sub{{font-size:14px;color:#2C3E50;line-height:1.75;
    max-width:500px;margin:0 auto 12px;}}
.fred-hero-flex{{font-size:14px;color:#2C3E50;line-height:1.6;
    max-width:460px;margin:0 auto 28px;}}
.fred-hero-flex span{{color:{BRAND_BLUE};font-weight:500;}}
.fred-hero-cta{{display:flex;flex-direction:column;align-items:center;gap:5px;}}
.fred-btn-primary{{background:{BRAND_BLUE};color:white;border:none;
    padding:13px 36px;border-radius:8px;font-size:15px;font-weight:500;
    cursor:pointer;display:flex;align-items:center;gap:9px;}}
.fred-btn-primary svg{{width:16px;height:16px;fill:white;}}
.fred-btn-reassure{{font-size:13px;color:{GREY};font-style:italic;margin:3px 0 6px;}}
.fred-btn-pricing{{font-size:13px;color:{GREY};}}
.fred-btn-pricing span{{color:{BRAND_BLUE};text-decoration:underline;
    text-underline-offset:3px;cursor:pointer;}}

.fred-divider{{border:none;border-top:0.5px solid #D5D8DC;margin:0;}}

.fred-section{{padding:44px 0;}}
.fred-sec-label{{font-size:11px;letter-spacing:2px;text-transform:uppercase;
    color:{GREY};margin-bottom:8px;text-align:center;}}
.fred-sec-title{{font-size:20px;font-weight:500;color:#1A252F;
    margin-bottom:16px;text-align:center;}}
.fred-sec-sub{{font-size:14px;color:#2C3E50;line-height:1.7;
    margin-bottom:20px;text-align:center;max-width:480px;
    margin-left:auto;margin-right:auto;}}

.fred-bullets{{list-style:none;display:flex;flex-direction:column;
    align-items:center;gap:10px;margin-bottom:24px;padding:0;}}
.fred-bullets li{{font-size:14px;color:#2C3E50;line-height:1.65;
    display:flex;align-items:flex-start;gap:9px;max-width:460px;}}
.fred-bdot{{width:6px;height:6px;min-width:6px;border-radius:50%;
    background:{BRAND_BLUE};margin-top:7px;}}
.fred-sub-bullet{{color:{BRAND_BLUE};font-weight:500;}}

.fred-upload-wrap{{display:flex;flex-direction:column;align-items:center;gap:8px;}}
.fred-upload-note{{font-size:11px;color:{GREY};font-style:italic;}}

.fred-traffic-legend{{background:#F4F6F7;border-radius:10px;
    padding:18px 20px;margin-bottom:16px;}}
.fred-traffic-title{{font-size:13px;font-weight:500;color:#1A252F;margin-bottom:12px;}}
.fred-trow{{display:flex;gap:11px;align-items:flex-start;margin-bottom:10px;}}
.fred-trow:last-child{{margin-bottom:0;}}
.fred-tdot{{width:13px;height:13px;min-width:13px;border-radius:50%;margin-top:3px;}}
.tdot-red{{background:{RED};}}
.tdot-amber{{background:{AMBER};}}
.tdot-green{{background:{GREEN};}}
.fred-ttext{{font-size:13px;color:#2C3E50;line-height:1.55;}}
.fred-ttext strong{{color:#1A252F;font-weight:500;}}

.fred-upload-zone{{border:1.5px dashed #BDC3C7;border-radius:10px;
    padding:24px;text-align:center;margin-bottom:10px;}}
.fred-upload-zone-title{{font-size:14px;font-weight:500;color:#1A252F;margin-bottom:4px;}}
.fred-upload-zone-sub{{font-size:12px;color:{GREY};margin-bottom:14px;}}
.fred-upload-tip{{background:#F4F6F7;border-radius:6px;padding:10px 14px;
    font-size:12px;color:{GREY};margin-top:6px;line-height:1.6;}}
.fred-upload-optional{{font-size:12px;color:{GREY};text-align:center;margin-top:8px;}}
.fred-upload-optional-link{{color:{BRAND_BLUE};text-decoration:underline;cursor:pointer;}}

.fred-sneak-header{{background:{BRAND_BLUE};color:white;padding:10px 16px;
    border-radius:6px 6px 0 0;font-size:13px;font-weight:500;}}
.fred-sneak-body{{padding:14px 16px;background:white;border:1px solid #D5D8DC;
    border-top:none;}}
.fred-sneak-entry{{font-size:12px;color:{GREY};font-style:italic;
    margin-bottom:10px;line-height:1.5;}}
.fred-sneak-more{{background:#F4F6F7;padding:16px;text-align:center;
    border:1px solid #D5D8DC;border-top:none;border-radius:0 0 6px 6px;}}
.fred-sneak-count{{font-size:14px;font-weight:500;color:#1A252F;margin-bottom:4px;}}
.fred-sneak-sub{{font-size:12px;color:{GREY};margin-bottom:8px;
    line-height:1.6;max-width:360px;margin-left:auto;margin-right:auto;}}
.fred-sneak-ready{{font-size:13px;font-weight:500;color:{BRAND_BLUE};margin-bottom:12px;}}

.fred-pricing-grid{{display:grid;
    grid-template-columns:repeat(auto-fit,minmax(190px,1fr));
    gap:14px;margin-top:4px;}}
.fred-price-card{{background:white;border:0.5px solid #D5D8DC;
    border-radius:12px;padding:20px;display:flex;flex-direction:column;}}
.fred-price-card.featured{{border:2px solid {BRAND_BLUE};}}
.fred-price-badge{{font-size:11px;font-weight:500;background:{BLUE_BG};
    color:#0C447C;padding:3px 10px;border-radius:6px;
    display:inline-block;margin-bottom:12px;}}
.fred-price-name{{font-size:14px;font-weight:500;color:#1A252F;margin-bottom:3px;}}
.fred-price-amount{{font-size:28px;font-weight:500;color:{BRAND_BLUE};margin-bottom:2px;}}
.fred-price-period{{font-size:11px;color:{GREY};margin-bottom:6px;line-height:1.5;}}
.fred-price-first{{font-size:11px;color:#2C3E50;background:#F4F6F7;
    border-radius:6px;padding:6px 10px;margin-bottom:12px;line-height:1.5;}}
.fred-price-features{{list-style:none;padding:0;display:flex;flex-direction:column;
    gap:7px;margin-bottom:18px;flex:1;}}
.fred-price-features li{{font-size:12px;color:#2C3E50;display:flex;gap:7px;
    align-items:flex-start;line-height:1.4;}}
.fred-price-features li::before{{content:"";width:5px;height:5px;min-width:5px;
    border-radius:50%;background:{BRAND_BLUE};margin-top:4px;}}

.fred-quote-box{{background:#F4F6F7;border-radius:12px;padding:28px 24px;}}
.fred-quote{{font-size:16px;color:#1A252F;line-height:1.75;
    font-style:italic;margin-bottom:14px;}}
.fred-quote-attr{{font-size:13px;color:{GREY};}}

.fred-faq-item{{padding:14px 0;border-bottom:0.5px solid #D5D8DC;}}
.fred-faq-item:last-child{{border-bottom:none;}}
.fred-faq-q{{font-size:14px;font-weight:500;color:#1A252F;margin-bottom:5px;}}
.fred-faq-a{{font-size:13px;color:#2C3E50;line-height:1.65;}}

.fred-footer{{padding:28px 0;border-top:0.5px solid #D5D8DC;text-align:center;}}
.fred-footer-logo{{font-size:16px;font-weight:500;color:{BRAND_BLUE};
    letter-spacing:3px;margin-bottom:8px;}}
.fred-footer-text{{font-size:11px;color:{GREY};line-height:1.8;}}

.fred-header-bar{{background:linear-gradient(135deg,{BRAND_BLUE},{BRAND_MID});
    color:white;padding:24px;border-radius:10px;margin-bottom:12px;}}
.fred-header-title{{font-size:40px;font-weight:500;letter-spacing:4px;margin:0;}}
.fred-header-sub{{font-size:13px;opacity:0.85;margin:4px 0 0 0;}}

.fred-beta-notice{{background:#FEF9E7;border-left:4px solid #F39C12;
    padding:12px 16px;border-radius:4px;font-size:13px;
    color:#7D6608;margin-bottom:20px;}}

.unlawful-flag{{border-left:4px solid {RED};padding:8px 12px;margin:6px 0;
    background:{RED_BG};border-radius:0 4px 4px 0;font-size:13px;
    color:#922B21;line-height:1.5;}}
.bestpractice-flag{{border-left:4px solid {AMBER};padding:8px 12px;margin:6px 0;
    background:{AMBER_BG};border-radius:0 4px 4px 0;font-size:13px;
    color:#7D6608;line-height:1.5;}}
.compliant-flag{{border-left:4px solid {GREEN};padding:8px 12px;margin:6px 0;
    background:{GREEN_BG};border-radius:0 4px 4px 0;font-size:13px;
    color:#1D6A36;line-height:1.5;}}
.pattern-flag{{border-left:4px solid {PURPLE};padding:8px 12px;margin:6px 0;
    background:{PURPLE_BG};border-radius:0 4px 4px 0;font-size:13px;
    color:#6C3483;line-height:1.5;}}
.tactical-flag{{border-left:4px solid {BRAND_BLUE};padding:8px 12px;margin:6px 0;
    background:{BLUE_BG};border-radius:0 4px 4px 0;font-size:13px;
    color:#1A3A5C;line-height:1.5;}}
.contradiction-flag{{border-left:4px solid {RED};padding:10px 14px;margin:8px 0;
    background:{RED_BG};border-radius:0 6px 6px 0;font-size:13px;
    color:#922B21;line-height:1.6;}}
.audit-header-red{{background:{RED};color:white;padding:10px 16px;
    border-radius:6px 6px 0 0;font-weight:500;font-size:13px;}}
.audit-header-amber{{background:{AMBER};color:white;padding:10px 16px;
    border-radius:6px 6px 0 0;font-weight:500;font-size:13px;}}
.audit-header-green{{background:{GREEN};color:white;padding:10px 16px;
    border-radius:6px 6px 0 0;font-weight:500;font-size:13px;}}
.audit-body{{background:white;border:1px solid #D5D8DC;border-top:none;
    padding:16px;border-radius:0 0 6px 6px;font-size:13px;
    line-height:1.7;margin-bottom:16px;}}
.anchor-line{{background:{BRAND_BLUE};color:white;padding:11px 16px;
    border-radius:6px;font-style:italic;font-size:13px;
    margin-top:10px;text-align:center;}}
.evidence-line{{background:#2C3E50;color:white;padding:9px 16px;
    border-radius:6px;font-style:italic;font-size:13px;
    margin-top:6px;text-align:center;}}
.review-capture{{background:{BLUE_BG};border:1px solid #AED6F1;
    border-radius:8px;padding:16px 20px;margin:16px 0;
    font-size:13px;color:#1A3A5C;line-height:1.6;}}
.subscription-signal{{background:linear-gradient(135deg,{BRAND_BLUE},{BRAND_MID});
    color:white;padding:22px 24px;border-radius:8px;
    margin:24px 0;font-size:14px;line-height:1.75;}}

.stButton>button{{background:{BRAND_BLUE};color:white;border:none;
    padding:10px 28px;border-radius:6px;font-weight:500;
    font-size:15px;width:100%;}}
.stButton>button:hover{{background:{BRAND_MID};}}
</style>
""", unsafe_allow_html=True)

# ─────────────────────────────────────────────
# SESSION STATE
# ─────────────────────────────────────────────

defaults = {
    'stage': 'landing',
    'answers': {},
    'extracted_sections': {},
    'report_results': [],
    'section_e_results': [],
    'policy_text': '',
    'raw_text': '',
    'email_text': '',
    'transcript_text': '',
    'correspondence_analysis': None,
    'post_meeting_email': '',
    'sneak_peek_result': None,
    'email_captured': False,
}
for key, val in defaults.items():
    if key not in st.session_state:
        st.session_state[key] = val

# ─────────────────────────────────────────────
# OFSTED STABLE PRINCIPLES
# ─────────────────────────────────────────────

OFSTED_PRINCIPLES = [
    {'area': 'Quality of education', 'principle': 'Ofsted inspection frameworks consistently expect schools to demonstrate that SEND pupils access a curriculum that is ambitious and appropriately adapted to their needs. Provision that lacks specificity makes this difficult to evidence at inspection.'},
    {'area': 'Leadership and management', 'principle': 'Schools are expected to demonstrate that leaders and managers have clear oversight of SEND provision and its effectiveness. An absence of delivery logs and monitoring records weakens this evidence base significantly.'},
    {'area': 'Personal development', 'principle': 'Inspection frameworks expect schools to show how SEND pupils are supported to develop confidence, resilience, and independence. Provision contingent on the child self-identifying need is unlikely to meet this expectation.'},
    {'area': 'Safeguarding', 'principle': 'Effective safeguarding requires that schools have specific, documented arrangements for pupils with identified vulnerabilities. Vague or discretionary provision creates risk that safeguarding responsibilities cannot be evidenced.'},
]

# ─────────────────────────────────────────────
# DOCUMENT READING
# ─────────────────────────────────────────────

def read_file(uploaded_file):
    if uploaded_file is None:
        return None, None
    name = uploaded_file.name.lower()
    if name.endswith('.pdf'):
        try:
            pdf_bytes = uploaded_file.read()
            doc = fitz.open(stream=pdf_bytes, filetype="pdf")
            text = "".join(page.get_text() for page in doc)
            doc.close()
            if len(text.strip()) < 50:
                return None, "This PDF appears to be image-based and could not be read as text. Try printing it to PDF from the original application."
            return text, None
        except Exception:
            return None, "This PDF could not be read. If it is password protected, open it, select print, and save as PDF — this removes the lock on most LA documents."
    elif name.endswith('.docx') or name.endswith('.doc'):
        try:
            doc = DocxDocument(uploaded_file)
            text = "\n".join(p.text for p in doc.paragraphs)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += "\n" + cell.text
            if len(text.strip()) < 50:
                return None, "This Word document appears to be empty or could not be read."
            return text, None
        except Exception:
            return None, "This Word document could not be read. If password protected, open it in Word, print to PDF, and upload the PDF instead."
    else:
        return None, "Format not supported. Please upload a PDF or Word document (.docx)."

def identify_sections(text):
    sections = {}
    patterns = {
        'A': r'(?:SECTION\s+A|Section\s+A)[:\s\-–—]*[^\n]*\n(.*?)(?=(?:SECTION\s+[B-K]|Section\s+[B-K])|$)',
        'B': r'(?:SECTION\s+B|Section\s+B)[:\s\-–—]*[^\n]*\n(.*?)(?=(?:SECTION\s+[C-K]|Section\s+[C-K])|$)',
        'C': r'(?:SECTION\s+C|Section\s+C)[:\s\-–—]*[^\n]*\n(.*?)(?=(?:SECTION\s+[D-K]|Section\s+[D-K])|$)',
        'D': r'(?:SECTION\s+D|Section\s+D)[:\s\-–—]*[^\n]*\n(.*?)(?=(?:SECTION\s+[E-K]|Section\s+[E-K])|$)',
        'E': r'(?:SECTION\s+E|Section\s+E)[:\s\-–—]*[^\n]*\n(.*?)(?=(?:SECTION\s+[F-K]|Section\s+[F-K])|$)',
        'F': r'(?:SECTION\s+F|Section\s+F)[:\s\-–—]*[^\n]*\n(.*?)(?=(?:SECTION\s+[G-K]|Section\s+[G-K])|$)',
        'G': r'(?:SECTION\s+G|Section\s+G)[:\s\-–—]*[^\n]*\n(.*?)(?=(?:SECTION\s+[H-K]|Section\s+[H-K])|$)',
        'H': r'(?:SECTION\s+H|Section\s+H)[:\s\-–—]*[^\n]*\n(.*?)(?=(?:SECTION\s+[I-K]|Section\s+[I-K])|$)',
        'I': r'(?:SECTION\s+I|Section\s+I)[:\s\-–—]*[^\n]*\n(.*?)(?=(?:SECTION\s+[J-K]|Section\s+[J-K])|$)',
        'J': r'(?:SECTION\s+J|Section\s+J)[:\s\-–—]*[^\n]*\n(.*?)(?=(?:SECTION\s+K|Section\s+K)|$)',
        'K': r'(?:SECTION\s+K|Section\s+K)[:\s\-–—]*[^\n]*\n(.*?)$',
    }
    for key, pattern in patterns.items():
        match = re.search(pattern, text, re.DOTALL | re.IGNORECASE)
        if match:
            content = re.sub(r'\n{3,}', '\n\n', match.group(1).strip())
            if len(content) > 20:
                sections[key] = content
    return sections

def extract_entries(text):
    numbered = re.split(r'\n\s*\d+[\.\)]\s+', text)
    if len(numbered) > 2:
        return [e.strip() for e in numbered if len(e.strip()) > 30]
    bulleted = re.split(r'\n\s*[\•\-\*]\s+', text)
    if len(bulleted) > 2:
        return [e.strip() for e in bulleted if len(e.strip()) > 30]
    paragraphs = re.split(r'\n{2,}', text)
    entries = [p.strip() for p in paragraphs if len(p.strip()) > 30]
    return entries if entries else [text]

def detect_doc_type(text):
    tl = text.lower()
    if any(kw in tl for kw in ['send policy', 'accessibility plan', 'behaviour policy', 'inclusion policy']):
        return 'policy'
    if any(kw in tl for kw in ['dear', 'kind regards', 'subject:', 'thank you for attending', 'thank you for coming']):
        return 'email'
    if any(kw in tl for kw in ['speaker', 'speaker 1', 'speaker 2', 'transcript', '[end]', 'yeah um', 'um ']):
        return 'transcript'
    if any(kw in tl for kw in ['section a', 'section b', 'section e', 'section f', 'education health and care']):
        return 'ehcp'
    if any(kw in tl for kw in ['educational psychologist', 'ep report', 'cognitive ability', 'standardised score']):
        return 'ep_report'
    if any(kw in tl for kw in ['occupational therapy', 'fine motor', 'sensory processing', 'ot report']):
        return 'ot_report'
    if any(kw in tl for kw in ['speech and language', 'salt report', 'communication assessment', 'language skills']):
        return 'salt_report'
    return 'other'

# ─────────────────────────────────────────────
# RULES ENGINE
# ─────────────────────────────────────────────

PROHIBITED = {
    r'\bshould\b': ('should', 'creates no lawful duty — it is a suggestion, not a statutory commitment'),
    r'\bcould\b': ('could', 'creates no lawful duty — possibility is not provision'),
    r'\bmay\b(?!\s+not)': ('may', 'means may not — no guaranteed entitlement is created'),
    r'\baccess to\b': ('access to', 'proximity to provision is not provision — no duty to deliver is created'),
    r'\bas needed\b': ('as needed', 'contingent on need being identified — who identifies it and how is unspecified'),
    r'\bwhere necessary\b': ('where necessary', 'entirely subjective — who determines necessity is unspecified'),
    r'\bas appropriate\b': ('as appropriate', 'discretionary — appropriate to whom and by what standard is left open'),
    r'\bregular\b': ('regular', 'unmeasurable — could mean daily, weekly, or termly without further specification'),
    r'\bencouraged\b': ('encouraged', 'creates no duty on any party — encouragement is not a statutory instruction'),
    r'\bmindful\b': ('mindful', 'an attitude, not a provision — no action is required under this wording'),
    r'\bcognisant\b': ('cognisant', 'awareness without obligation — no specified action or lawful duty'),
    r'\bholistic approach\b': ('holistic approach', 'undefined — no named strategy, approach, or measurable outcome'),
    r'\bopportunity\b': ('opportunity', 'possibility is not guaranteed provision under the Children and Families Act 2014'),
    r'\bit is expected\b': ('it is expected', 'expectation creates no lawful duty on any party'),
    r'\bwould benefit from\b': ('would benefit from', 'assessment language — must be converted to a specified commitment in Section F'),
    r'\bit is recommended\b': ('it is recommended', 'recommendation language — must be converted to a specified commitment in Section F'),
    r'\bat their.*?discretion\b': ('at their discretion', 'professional discretion cannot override a statutory entitlement'),
    r'\bwhere possible\b': ('where possible', 'conditional — possibility is not guaranteed provision'),
    r'\bas directed by\b': ('as directed by', 'places statutory provision under the daily discretion of another party'),
    r'\bflexib\w*\b': ('flexible/flexibility', 'unmeasurable — who decides what is flexible and when is unspecified'),
    r'\bresponsive\b': ('responsive', 'reactive delivery is not specified provision — frequency and trigger criteria must be stated'),
    r'\btailored\b': ('tailored', 'undefined without specifying what the tailoring consists of'),
    r'\bembedded\b': ('embedded', 'not a quantified description of delivery — frequency and context must be stated'),
}

UNIVERSAL = [
    'high-quality teaching', 'quality first teaching', 'broad and balanced curriculum',
    'differentiated curriculum', 'universal offer', 'graduated response',
    'scaffolding for tasks', 'quality teaching', 'ordinarily available',
]

QUANT = {
    'frequency': r'\b(\d+\s*(?:times?|sessions?|hours?)\s*(?:per|a|each)\s*(?:week|day|term|month)|daily|weekly|fortnightly|monthly|termly|once|twice)\b',
    'duration': r'\b(\d+\s*(?:minutes?|hours?|mins?))\b',
    'role': r'\b(therapist|psychologist|specialist|SENCO|senco|teacher|LSA|TA|teaching assistant|learning support|coordinator|practitioner|nurse|OT|SALT|SLT|occupational|speech|language)\b',
    'named_individual': r'\b(Mrs|Mr|Ms|Dr|Miss)\s+[A-Z][a-z]+\b',
}

def chk_quant(text):
    return {k: bool(re.search(p, text, re.IGNORECASE)) for k, p in QUANT.items()}

def chk_prohibited(text):
    findings = []
    seen = set()
    tl = text.lower()
    for pattern, (term, exp) in PROHIBITED.items():
        if term not in seen and re.search(pattern, tl, re.IGNORECASE):
            findings.append((term, exp))
            seen.add(term)
    return findings

def chk_universal(text):
    tl = text.lower()
    return [i for i in UNIVERSAL if i in tl]

def chk_laundering(text):
    patterns = [r'\bwould benefit from\b', r'\bit is recommended\b',
                r'\bit is suggested\b', r'\bmay benefit from\b']
    return [p.replace(r'\b', '') for p in patterns if re.search(p, text, re.IGNORECASE)]

def chk_dilution(text):
    patterns = [r'\bshared with other\b', r'\bmay be shared\b',
                r'\bas resources allow\b', r'\bsubject to availability\b',
                r'\bwhen staff are available\b', r'\bdepending on resources\b',
                r'\bwider class\b', r"\bat the school'?s discretion\b"]
    return [p.replace(r'\b', '') for p in patterns if re.search(p, text, re.IGNORECASE)]

def chk_policy(entry, policy):
    if not policy:
        return []
    gaps = []
    pl = policy.lower()
    el = entry.lower()
    commitments = [
        ('1:1 support', ['1:1', 'one to one', 'individual support']),
        ('named key worker', ['key worker', 'key person']),
        ('sensory assessment', ['sensory assessment', 'sensory profile', 'sensory audit']),
        ('home-school communication', ['parent update', 'home school']),
        ('risk assessment', ['risk assess']),
        ('accessibility arrangements', ['accessible', 'adaptations']),
    ]
    for label, keywords in commitments:
        if any(kw in pl for kw in keywords) and not any(kw in el for kw in keywords):
            gaps.append(
                f"The school's own policy references {label}. "
                f"This does not appear in this provision entry. "
                f"The school cannot dispute what its own policy commits to — "
                f"worth raising at annual review."
            )
    return gaps

def is_compliant(text, quant):
    has_must = bool(re.search(r'\bmust\b', text, re.IGNORECASE))
    prohibited = chk_prohibited(text)
    universal = chk_universal(text)
    return (has_must and quant.get('frequency') and quant.get('duration')
            and quant.get('role') and not prohibited and not universal)

def get_ofsted(text):
    tl = text.lower()
    if any(w in tl for w in ['safe', 'risk', 'physical', 'behaviour', 'incident']):
        return OFSTED_PRINCIPLES[3]
    if any(w in tl for w in ['independent', 'confidence', 'resilience']):
        return OFSTED_PRINCIPLES[2]
    if any(w in tl for w in ['monitor', 'oversight', 'review', 'log', 'record']):
        return OFSTED_PRINCIPLES[1]
    return OFSTED_PRINCIPLES[0]

def audit_entry(entry_text, entry_number, policy_text=''):
    quant = chk_quant(entry_text)
    prohibited = chk_prohibited(entry_text)
    universal = chk_universal(entry_text)
    laundering = chk_laundering(entry_text)
    dilution = chk_dilution(entry_text)
    compliant = is_compliant(entry_text, quant)
    policy_gaps = chk_policy(entry_text, policy_text)

    unlawful = []
    for term, exp in prohibited:
        unlawful.append(f'"{term}" — {exp}.')
    if not quant['frequency']:
        unlawful.append('No frequency specified — how often provision is delivered is not stated. The SEND Code of Practice requires provision to be specified and quantified. Without frequency this provision cannot be monitored or enforced.')
    if not quant['duration']:
        unlawful.append('No duration specified — the length of each session is not stated. Provision without quantification cannot be measured or challenged at annual review.')
    if not quant['role']:
        unlawful.append('No deliverer role specified — who provides this provision and at what qualification or training level is not stated. A lawful duty requires a named responsible role, not just a description of activity.')

    patterns = []
    if universal:
        patterns.append('Universal provision identified — this entry describes what the school is already required to provide all pupils. Its presence in Section F creates no additional lawful entitlement specific to this child. Section F must contain provision above and beyond the school\'s universal offer.')
    if laundering:
        patterns.append('Recommendation laundering identified — assessment or report language has been copied into Section F without being converted into a specified lawful commitment. Referencing the existence of professional advice without acting on it creates no enforceable duty under the Children and Families Act 2014.')
    if dilution:
        patterns.append('Dilution clause identified — wording allows this provision to be shared or made conditional on school resources or staffing. An individual statutory entitlement cannot be diluted at the school\'s discretion.')

    best_practice = []
    if not quant['named_individual']:
        best_practice.append(
            'No named accountable person — the lawful requirement is that the deliverer role '
            'and training level are specified. As a best practice consideration, naming the '
            'SENCO as the accountable person supports continuity and makes monitoring easier '
            'to evidence at annual review and at inspection. '
            'This is a wellbeing recommendation, not a lawful requirement.'
        )
    if not re.search(r'\b(review|reviewed|assess|monitor|evaluated)\b', entry_text, re.IGNORECASE):
        best_practice.append('No review mechanism stated — provision without a stated review mechanism cannot be assessed for effectiveness. Consider asking at annual review how the effectiveness of this provision is assessed and recorded.')

    required = []
    if not compliant:
        if not quant['frequency']:
            required.append('Frequency must be stated — number of sessions per week or per term, specified plainly')
        if not quant['duration']:
            required.append('Duration must be stated — length of each session in minutes')
        if not quant['role']:
            required.append('Deliverer role must be named — role title and relevant qualification or training level specified')
        if universal:
            required.append('Entry must describe provision additional to the universal offer — specific to this child\'s identified needs')
        if laundering:
            required.append('Professional recommendations must be reproduced as specified provision — not referenced as existing advice')
        if dilution:
            required.append('Shared or conditional wording must be removed — provision specified as an individual guaranteed entitlement')
        required.append(
            'Mandatory delivery log — all provision recorded in a dated delivery log '
            'showing date, duration, who delivered, and any relevant observations. '
            'This is the evidence base for the Do stage of the school\'s statutory '
            'APDR (Assess, Plan, Do, Review) cycle. Without it the Review stage '
            'cannot be conducted accurately and the cycle breaks down.'
        )

    tactical = [
        'Request the Physical Delivery Log for this provision. '
        'Dated entries must show each session — date, duration, who delivered, and format. '
        'If no log exists there is no evidence this provision has been delivered. '
        'Lack of evidence is evidence of lack.'
    ]
    if not compliant:
        tactical.append(
            'At your next annual review this entry must be rewritten to full specification standard. '
            'FRED will remind you of this finding as your review approaches — '
            'enter your review date at the bottom of this report.'
        )
    if dilution:
        tactical.append('Request written confirmation of how many other pupils share this provision and what proportion of the named support this child actually receives.')

    return {
        'entry_number': entry_number,
        'entry_text': entry_text,
        'is_compliant': compliant,
        'unlawful_deficiencies': unlawful,
        'additional_patterns': patterns,
        'best_practice_gaps': best_practice,
        'ofsted_principle': get_ofsted(entry_text),
        'policy_gaps': policy_gaps,
        'required_specification': required,
        'tactical_advice': tactical,
    }

def audit_section_e(text):
    results = []
    outcomes = re.split(r'\n\s*[\•\-\*\d][\.\)]?\s+', text)
    outcomes = [o.strip() for o in outcomes if len(o.strip()) > 20]
    if not outcomes:
        outcomes = [p.strip() for p in text.split('\n') if len(p.strip()) > 20]
    for i, outcome in enumerate(outcomes):
        ol = outcome.lower()
        unlawful = []
        bp = []
        if not re.search(r'\b(currently|baseline|starting point|at present|now)\b', ol):
            unlawful.append('No baseline stated — without a starting point progress cannot be objectively measured at annual review. The SEND Code of Practice requires outcomes to be measurable.')
        if not re.search(r'\b(\d+|percentage|score|level|times|independently|consistently|measured by|assessed)\b', ol):
            unlawful.append('No measurable indicator — success cannot be objectively assessed. An outcome without a measurable indicator cannot be reviewed under the APDR cycle.')
        if not re.search(r'\b(by|within|term|year|month|weeks?|annual review|end of)\b', ol):
            bp.append('No timeframe stated — when this outcome should be achieved is not specified. This supports effective APDR cycle review and annual review preparation.')
        results.append({'outcome_number': i+1, 'outcome_text': outcome, 'unlawful_failures': unlawful, 'best_practice_gaps': bp})
    return results

# ─────────────────────────────────────────────
# CORRESPONDENCE ENGINE
# ─────────────────────────────────────────────

UNENFORCEABLE_EMAIL = [
    ('in place', 'Claims provision is in place without referencing any delivery record'),
    ('regularly', '"Regularly" is unmeasurable — frequency must be stated'),
    ('as outlined', 'References the plan without evidencing delivery'),
    ('consistently', 'A claim — the delivery log is the evidence'),
    ('embedded', 'Not a quantified description of delivery'),
    ('responsive', 'Reactive delivery is not specified provision'),
    ('tailored to his needs', 'Undefined without specifying what the tailoring consists of'),
    ('some flexibility', 'Flexibility in EHCP provision is not permitted — provision must be delivered as specified'),
    ('monitor', 'Monitoring without a stated recording method is unverifiable'),
]

def analyse_correspondence(email_text, ehcp_sections, transcript_text=''):
    analysis = {
        'unenforceable_claims': [],
        'contradictions_with_transcript': [],
        'deflected_items': [],
        'addressed_items': [],
    }
    el = email_text.lower()

    for term, exp in UNENFORCEABLE_EMAIL:
        if term in el:
            analysis['unenforceable_claims'].append(
                f'"{term}" — {exp}. A delivery log is required to substantiate this claim. Lack of evidence is evidence of lack.'
            )

    if 'F' in ehcp_sections:
        fl = ehcp_sections['F'].lower()
        provisions = [
            ('social skills group', ['social skills', 'social group']),
            ('emotional regulation sessions', ['emotional regulation', 'emotional literacy']),
            ('sensory and movement breaks', ['sensory', 'movement break', 'calm space']),
            ('adult support', ['adult support', 'lsa', 'full-time support']),
            ('speech and language support', ['speech', 'language', 'salt', 'communication']),
            ('occupational therapy', ['fine motor', 'ot', 'occupational']),
        ]
        for label, keywords in provisions:
            in_ehcp = any(kw in fl for kw in keywords)
            in_email = any(kw in el for kw in keywords)
            if in_ehcp and not in_email:
                analysis['deflected_items'].append(
                    f'{label.title()} — specified in the EHCP but not addressed in this email. No confirmation of delivery has been provided.'
                )
            elif in_ehcp and in_email:
                analysis['addressed_items'].append(
                    f'{label.title()} — referenced in both the EHCP and this email. Request the delivery log to substantiate any claims of delivery.'
                )

    if transcript_text:
        tl = transcript_text.lower()
        if ('not every week' in tl or 'staffing' in tl) and 'social' in tl:
            if 'weekly' in el or 'in place' in el:
                analysis['contradictions_with_transcript'].append(
                    'Social skills group — the transcript records a direct admission that sessions have not taken place every week and that staffing has been a difficulty. The email states sessions are in place weekly. These two accounts are not consistent. Written clarification and the full delivery record are required.'
                )
        if "don't formally track" in tl or 'not formally track' in tl or 'staff just know' in tl:
            if 'sensory' in el and ('regular' in el or 'responsive' in el):
                analysis['contradictions_with_transcript'].append(
                    'Sensory breaks — the transcript records that provision is not formally tracked and that no delivery log exists. The email describes provision as regular and responsive. Without a delivery log these claims cannot be evidenced. Lack of evidence is evidence of lack.'
                )
        if 'someone else' in tl or 'if not me' in tl:
            if 'full-time support' in el or 'adult support' in el:
                analysis['contradictions_with_transcript'].append(
                    'Adult support — the transcript records that support is provided by different adults across the day without a named accountable person. The email presents full-time support as a consistent guaranteed provision. Please confirm in writing the roles, training, and accountability arrangements for all adults providing support.'
                )
        if 'mostly mornings' in tl or ('mornings' in tl and 'transitions' in tl):
            if 'consistently across' in el or 'embedded consistently' in el:
                analysis['contradictions_with_transcript'].append(
                    'Visual supports — the transcript describes use as mostly during mornings and transitions. The email states these are embedded consistently across the day. Please clarify which is accurate and provide the delivery record.'
                )
    return analysis

def generate_post_meeting_email(analysis, answers):
    tone = answers.get('q5', 'Constructive but cautious')
    openings = {
        'Warm and collaborative': 'Thank you for the meeting and for your follow up email. We found the discussion helpful and want to ensure our understanding of what was discussed is accurately recorded.',
        'Constructive but cautious': 'Thank you for the meeting and for your follow up email. We write to ensure our understanding of what was discussed is accurately recorded.',
        'Professionally firm': 'Thank you for your follow up email. We write to record our understanding of the meeting, which differs in some respects from the summary you have provided.',
        'Formally assertive': 'We write further to the recent meeting and your subsequent email. We wish to place on record our understanding of what was agreed and what remains unresolved.',
        'Rights-based and formal': 'We write further to the meeting and your subsequent correspondence. The following sets out our understanding of what was discussed and what outstanding matters require a written response.',
    }
    parts = [openings.get(tone, openings['Constructive but cautious']), '']
    if analysis['contradictions_with_transcript'] or analysis['deflected_items']:
        parts.append('What requires a written response\n')
        for c in analysis['contradictions_with_transcript']:
            parts.append(c + '\n')
        for d in analysis['deflected_items']:
            parts.append(d + '\n')
    if analysis['addressed_items']:
        parts.append('Provision referenced — delivery log requested\n')
        for a in analysis['addressed_items']:
            parts.append(f'— {a.split(" — ")[0]}')
    parts.append('')
    parts.append(
        'Please let us know within five working days if anything above does not reflect '
        'your understanding of the meeting. If we do not hear from you within that time '
        'we will treat this summary as the agreed record.'
    )
    return '\n'.join(parts)

# ─────────────────────────────────────────────
# RENDER FUNCTIONS
# ─────────────────────────────────────────────

def render_traffic_legend():
    st.markdown(f"""
    <div class="fred-traffic-legend">
        <div class="fred-traffic-title">Here is how FRED colour codes its findings</div>
        <div class="fred-trow">
            <div class="fred-tdot tdot-red"></div>
            <div class="fred-ttext">
                <strong>Red — lawful requirement not met.</strong>
                The provision does not meet the statutory standard set by the
                Children and Families Act 2014. Must be addressed at annual review.
            </div>
        </div>
        <div class="fred-trow">
            <div class="fred-tdot tdot-amber"></div>
            <div class="fred-ttext">
                <strong>Amber — best practice gap.</strong>
                Meets the minimum lawful standard but falls short of what good
                practice recommends. Worth raising at annual review.
            </div>
        </div>
        <div class="fred-trow">
            <div class="fred-tdot tdot-green"></div>
            <div class="fred-ttext">
                <strong>Green — compliant.</strong>
                Meets the lawful standard. Use compliant entries as the
                benchmark when challenging non-compliant ones.
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)

def render_sneak_peek(result):
    entry_preview = result['entry_text'][:200]
    unlawful = result['unlawful_deficiencies'][:3]
    count_remaining = max(0, (len(st.session_state.report_results) - 1))

    st.markdown(f"""
    <div class="fred-sneak-header">FRED has read your plan — here is one finding</div>
    <div class="fred-sneak-body">
        <div class="fred-sneak-entry">"{entry_preview}{'...' if len(result['entry_text']) > 200 else ''}"</div>
        {''.join(f'<div class="unlawful-flag">⚠ {d}</div>' for d in unlawful)}
        <div class="anchor-line">If it is not specified and evidenced, it is not lawfully enforceable under the Children and Families Act 2014.</div>
        <div class="evidence-line">Lack of evidence is evidence of lack.</div>
    </div>
    <div class="fred-sneak-more">
        <div class="fred-sneak-count">This is one entry from Section F of your plan</div>
        <div class="fred-sneak-sub">Your full report covers every provision entry across Section F and Section E outcomes — with tactical advice and required specification for each finding.</div>
        <div class="fred-sneak-ready">Your report is ready.</div>
    </div>
    """, unsafe_allow_html=True)

def render_correspondence(analysis, post_meeting_email):
    st.markdown("## Correspondence analysis")
    if analysis['contradictions_with_transcript']:
        st.markdown("### Contradictions — transcript vs email")
        st.markdown("*The following claims in the email are not consistent with what the transcript records.*")
        for c in analysis['contradictions_with_transcript']:
            st.markdown(f'<div class="contradiction-flag">⚠ {c}</div>', unsafe_allow_html=True)
        st.markdown('<div class="evidence-line">Lack of evidence is evidence of lack.</div>', unsafe_allow_html=True)
    if analysis['unenforceable_claims']:
        st.markdown("### Unsubstantiated claims in email")
        for u in analysis['unenforceable_claims']:
            st.markdown(f'<div class="unlawful-flag">⚠ {u}</div>', unsafe_allow_html=True)
    if analysis['deflected_items']:
        st.markdown("### Provision not addressed in email")
        for d in analysis['deflected_items']:
            st.markdown(f'<div class="bestpractice-flag">◉ {d}</div>', unsafe_allow_html=True)
    if analysis['addressed_items']:
        st.markdown("### Provision referenced — delivery log required")
        for a in analysis['addressed_items']:
            st.markdown(f'<div class="tactical-flag">→ {a}</div>', unsafe_allow_html=True)
    st.markdown("---")
    st.markdown("### Post-meeting summary email")
    st.markdown("*Send this within 24 hours. The school has five working days to correct anything. Silence is acceptance.*")
    st.text_area("Copy and send:", value=post_meeting_email, height=360, key="post_meeting_output")

def render_full_report(report_results, section_e_results, answers):
    st.markdown("---")
    st.markdown("## Your FRED report")

    ehcp_status = answers.get('q2', 'Unknown')
    process_stage = answers.get('q3', 'Not specified')

    st.markdown(f"""
    <div style="background:#F4F6F7;border-radius:8px;padding:13px 17px;margin:10px 0;font-size:13px;">
        <strong>Status:</strong> {ehcp_status} &nbsp;|&nbsp;
        <strong>Stage:</strong> {process_stage}
    </div>
    """, unsafe_allow_html=True)

    if 'final' in ehcp_status.lower():
        st.warning(
            "**Final EHCP pathway active.** This plan has been formally issued by the LA. "
            "The school is now responsible for delivery. All findings below inform what you "
            "raise at annual review — not changes to the current document."
        )

    render_traffic_legend()

    if section_e_results:
        st.markdown("### Section E — Outcomes")
        for r in section_e_results:
            has_issues = r['unlawful_failures'] or r['best_practice_gaps']
            if not has_issues:
                st.markdown(f"""
                <div class="audit-header-green">Outcome {r['outcome_number']} — compliant</div>
                <div class="audit-body">
                    <div class="compliant-flag">✓ This outcome meets the SMART standard.</div>
                    <em>"{r['outcome_text'][:200]}"</em>
                </div>""", unsafe_allow_html=True)
            else:
                st.markdown(f"""
                <div class="audit-header-red">Outcome {r['outcome_number']} — review required</div>
                <div class="audit-body">
                    <em>"{r['outcome_text'][:200]}"</em><br><br>
                    {''.join(f'<div class="unlawful-flag">⚠ {f}</div>' for f in r['unlawful_failures'])}
                    {''.join(f'<div class="bestpractice-flag">◉ {g}</div>' for g in r['best_practice_gaps'])}
                </div>""", unsafe_allow_html=True)
        st.markdown("---")

    if report_results:
        st.markdown("### Section F — Provision")
        unlawful_count = sum(1 for r in report_results if r['unlawful_deficiencies'] or r['additional_patterns'])
        compliant_count = sum(1 for r in report_results if r['is_compliant'])
        total = len(report_results)

        c1, c2, c3 = st.columns(3)
        c1.metric("Total entries", total)
        c2.metric("Lawful requirement not met", unlawful_count,
                 delta=f"{unlawful_count} entries" if unlawful_count > 0 else None,
                 delta_color="inverse")
        c3.metric("Compliant", compliant_count)
        st.markdown("<br>", unsafe_allow_html=True)

        for result in report_results:
            if result['is_compliant']:
                st.markdown(f"""
                <div class="audit-header-green">Provision {result['entry_number']} — compliant</div>
                <div class="audit-body">
                    <div class="compliant-flag">✓ Meets the lawful specification standard. Use as benchmark at annual review.</div>
                    <em>"{result['entry_text'][:300]}"</em>
                </div>""", unsafe_allow_html=True)
            else:
                has_unlawful = bool(result['unlawful_deficiencies'] or result['additional_patterns'])
                hclass = 'audit-header-red' if has_unlawful else 'audit-header-amber'
                hlabel = 'lawful requirement not met' if has_unlawful else 'best practice gap'
                st.markdown(f"""
                <div class="{hclass}">Provision {result['entry_number']} — {hlabel}</div>
                <div class="audit-body">
                <em>"{result['entry_text'][:300]}{'...' if len(result['entry_text']) > 300 else ''}"</em><br><br>
                """, unsafe_allow_html=True)
                if result['unlawful_deficiencies']:
                    st.markdown("**Lawful requirements not met**")
                    for d in result['unlawful_deficiencies']:
                        st.markdown(f'<div class="unlawful-flag">⚠ {d}</div>', unsafe_allow_html=True)
                if result['additional_patterns']:
                    st.markdown("**Additional pattern identified**")
                    for p in result['additional_patterns']:
                        st.markdown(f'<div class="pattern-flag">◈ {p}</div>', unsafe_allow_html=True)
                if result['best_practice_gaps']:
                    st.markdown("**Best practice gaps**")
                    for g in result['best_practice_gaps']:
                        st.markdown(f'<div class="bestpractice-flag">◉ {g}</div>', unsafe_allow_html=True)
                if result['ofsted_principle']:
                    op = result['ofsted_principle']
                    st.markdown("**Inspection framework note**")
                    st.markdown(f'<div class="bestpractice-flag"><strong>{op["area"]}:</strong> {op["principle"]}</div>', unsafe_allow_html=True)
                if result['policy_gaps']:
                    st.markdown("**School policy cross-reference**")
                    for pg in result['policy_gaps']:
                        st.markdown(f'<div class="pattern-flag">◈ {pg}</div>', unsafe_allow_html=True)
                if result['required_specification']:
                    st.markdown("**Required specification**")
                    for spec in result['required_specification']:
                        st.markdown(f"— {spec}")
                if result['tactical_advice']:
                    st.markdown("**Tactical advice**")
                    for advice in result['tactical_advice']:
                        st.markdown(f'<div class="tactical-flag">→ {advice}</div>', unsafe_allow_html=True)
                if result['unlawful_deficiencies']:
                    st.markdown("""
                    <div class="anchor-line">If it is not specified and evidenced, it is not lawfully enforceable under the Children and Families Act 2014.</div>
                    <div class="evidence-line">Lack of evidence is evidence of lack.</div>
                    """, unsafe_allow_html=True)
                st.markdown("</div><br>", unsafe_allow_html=True)

        st.info("Upload the expert reports (EP, OT, or SLT) to begin the Cross-Reference report.")

    st.markdown(f"""
    <div class="review-capture">
        <strong>Hold this for your annual review.</strong><br>
        Enter your next annual review date and FRED will begin working through
        these findings with you in the weeks before it. Nothing will be forgotten.
    </div>
    """, unsafe_allow_html=True)
    st.date_input("Annual review date (optional):", key="review_date")

    unlawful_total = sum(
        len(r['unlawful_deficiencies']) + len(r['additional_patterns'])
        for r in report_results
    )
    if unlawful_total > 0:
        st.markdown(f"""
        <div class="subscription-signal">
            <strong>FRED has identified {unlawful_total} provision failures in this plan.</strong><br><br>
            The full FRED service will hold these findings, track whether the school
            delivers on its obligations, draft your correspondence, prepare you for
            the annual review meeting with a script you can read in the room, and
            produce the post-meeting summary that puts everything on the written record.<br><br>
            Annual subscription — from £XX per year.
            Less than the cost of a single hour with a specialist advocate.
        </div>
        """, unsafe_allow_html=True)

# ─────────────────────────────────────────────
# DOCUMENT GENERATION
# ─────────────────────────────────────────────

def generate_docx(report_results, section_e_results, answers):
    doc = DocxDocument()
    RED_C = RGBColor(0xC0, 0x39, 0x2B)
    AMBER_C = RGBColor(0xD4, 0xA0, 0x17)
    GREEN_C = RGBColor(0x1E, 0x84, 0x49)
    BLUE_C = RGBColor(0x1B, 0x4F, 0x72)
    PURPLE_C = RGBColor(0x8E, 0x44, 0xAD)

    def h(text, level=1, c=RGBColor(0x1B, 0x4F, 0x72)):
        heading = doc.add_heading(text, level=level)
        if heading.runs:
            heading.runs[0].font.color.rgb = c
        return heading

    def p(text, c=RGBColor(0,0,0), size=10, bold=False, italic=False):
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.font.color.rgb = c
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        return para

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("FRED")
    r.font.size = Pt(36); r.font.bold = True; r.font.color.rgb = BLUE_C

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.add_run("Families' Rights and Entitlements Directory — EHCP Report").font.color.rgb = RGBColor(0x2E, 0x86, 0xC1)
    doc.add_paragraph(f"Status: {answers.get('q2','Unknown')} | Beta v0.5")
    doc.add_paragraph("FRED provides information to help you understand the language of your child's plan and what the law says about it. It does not constitute legal advice.")
    doc.add_page_break()

    h("Output key", level=2)
    p("● Red — lawful requirement not met. Must be addressed.", RED_C)
    p("● Amber — best practice gap. Recommended.", AMBER_C)
    p("● Green — compliant. Meets the lawful standard.", GREEN_C)
    doc.add_paragraph()

    if section_e_results:
        h("Section E — Outcomes")
        for r_ in section_e_results:
            c_ = RED_C if r_['unlawful_failures'] else (AMBER_C if r_['best_practice_gaps'] else GREEN_C)
            h(f"Outcome {r_['outcome_number']}", level=2, c=c_)
            p(f'"{r_["outcome_text"]}"', italic=True)
            for f_ in r_['unlawful_failures']:
                p(f"⚠ {f_}", RED_C)
            for g_ in r_['best_practice_gaps']:
                p(f"◉ {g_}", AMBER_C)
            if not r_['unlawful_failures'] and not r_['best_practice_gaps']:
                p("✓ Meets SMART criteria.", GREEN_C)
        doc.add_page_break()

    if report_results:
        h("Section F — Provision")
        for result in report_results:
            c_ = (GREEN_C if result['is_compliant']
                 else RED_C if result['unlawful_deficiencies']
                 else AMBER_C)
            label = ("Compliant" if result['is_compliant']
                    else "Lawful requirement not met" if result['unlawful_deficiencies']
                    else "Best practice gap")
            h(f"Provision {result['entry_number']} — {label}", level=2, c=c_)
            p(f'"{result["entry_text"][:400]}"', italic=True)
            if result['unlawful_deficiencies']:
                h("Lawful requirements not met", level=3, c=RED_C)
                for d in result['unlawful_deficiencies']:
                    p(f"⚠ {d}", RED_C)
            if result['additional_patterns']:
                h("Additional pattern identified", level=3, c=PURPLE_C)
                for pt in result['additional_patterns']:
                    p(f"◈ {pt}", PURPLE_C)
            if result['best_practice_gaps']:
                h("Best practice gaps", level=3, c=AMBER_C)
                for g in result['best_practice_gaps']:
                    p(f"◉ {g}", AMBER_C)
            if result['ofsted_principle']:
                op = result['ofsted_principle']
                h("Inspection framework note", level=3, c=AMBER_C)
                p(f"{op['area']}: {op['principle']}", AMBER_C)
            if result['policy_gaps']:
                h("School policy cross-reference", level=3, c=PURPLE_C)
                for pg in result['policy_gaps']:
                    p(f"◈ {pg}", PURPLE_C)
            if result['required_specification']:
                h("Required specification", level=3)
                for spec in result['required_specification']:
                    doc.add_paragraph(spec, style='List Bullet')
            if result['tactical_advice']:
                h("Tactical advice", level=3, c=BLUE_C)
                for advice in result['tactical_advice']:
                    p(f"→ {advice}", BLUE_C)
            if result['unlawful_deficiencies']:
                p("If it is not specified and evidenced, it is not lawfully enforceable under the Children and Families Act 2014. Lack of evidence is evidence of lack.", BLUE_C, bold=True, italic=True)
            doc.add_paragraph()

    doc.add_page_break()
    h("About the full FRED service")
    doc.add_paragraph("The full FRED service holds all your documents, tracks your correspondence, drafts emails calibrated to your relationship with the school, prepares you for every meeting with a script and agenda, and stays with you through every annual review.")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def generate_pdf(report_results, section_e_results, answers):
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                           rightMargin=20*mm, leftMargin=20*mm,
                           topMargin=20*mm, bottomMargin=20*mm)
    styles = getSampleStyleSheet()
    brand = HexColor('#1B4F72')
    red = HexColor('#C0392B')
    amber = HexColor('#D4A017')
    green = HexColor('#1E8449')
    purple = HexColor('#8E44AD')

    def ps(name, parent='Normal', **kw):
        return ParagraphStyle(name, parent=styles[parent], **kw)

    h1 = ps('H1', 'Heading1', textColor=brand, fontSize=16, spaceAfter=6)
    h2r = ps('H2R', 'Heading2', textColor=red, fontSize=13, spaceAfter=4)
    h2a = ps('H2A', 'Heading2', textColor=amber, fontSize=13, spaceAfter=4)
    h2g = ps('H2G', 'Heading2', textColor=green, fontSize=13, spaceAfter=4)
    h3 = ps('H3', 'Heading3', fontSize=11, spaceAfter=4)
    body = ps('Body', fontSize=10, spaceAfter=4, leading=15)
    red_s = ps('RS', fontSize=10, textColor=red, leftIndent=10, spaceAfter=3, leading=14)
    amb_s = ps('AS', fontSize=10, textColor=amber, leftIndent=10, spaceAfter=3, leading=14)
    grn_s = ps('GS', fontSize=10, textColor=green, leftIndent=10, spaceAfter=3, leading=14)
    pur_s = ps('PS', fontSize=10, textColor=purple, leftIndent=10, spaceAfter=3, leading=14)
    tac_s = ps('TS', fontSize=10, textColor=brand, leftIndent=10, spaceAfter=3, leading=14)
    anc_s = ps('ANS', fontSize=10, textColor=brand, fontName='Helvetica-BoldOblique', spaceAfter=8, leading=14)

    story = []
    story.append(Paragraph("FRED", ps('TT', 'Title', textColor=brand, fontSize=32)))
    story.append(Paragraph("Families' Rights and Entitlements Directory", h1))
    story.append(Spacer(1, 5*mm))
    story.append(Paragraph(f"EHCP Report | Status: {answers.get('q2','Unknown')} | Beta v0.5", body))
    story.append(Paragraph("This report provides information to help you understand the language of your child's plan and what the law says about it. It does not constitute legal advice.", body))
    story.append(Spacer(1, 4*mm))
    story.append(Paragraph("Output key", h1))
    story.append(Paragraph("● Red — lawful requirement not met. Must be addressed.", red_s))
    story.append(Paragraph("● Amber — best practice gap. Recommended.", amb_s))
    story.append(Paragraph("● Green — compliant. Meets the lawful standard.", grn_s))
    story.append(Spacer(1, 5*mm))

    if section_e_results:
        story.append(Paragraph("Section E — Outcomes", h1))
        for r_ in section_e_results:
            h_ = h2r if r_['unlawful_failures'] else (h2a if r_['best_practice_gaps'] else h2g)
            story.append(Paragraph(f"Outcome {r_['outcome_number']}", h_))
            story.append(Paragraph(f'<i>"{r_["outcome_text"][:300]}"</i>', body))
            for f_ in r_['unlawful_failures']:
                story.append(Paragraph(f"⚠ {f_}", red_s))
            for g_ in r_['best_practice_gaps']:
                story.append(Paragraph(f"◉ {g_}", amb_s))
            if not r_['unlawful_failures'] and not r_['best_practice_gaps']:
                story.append(Paragraph("✓ Meets SMART criteria.", grn_s))
            story.append(Spacer(1, 3*mm))

    if report_results:
        story.append(Paragraph("Section F — Provision", h1))
        for result in report_results:
            h_ = (h2g if result['is_compliant']
                 else h2r if result['unlawful_deficiencies']
                 else h2a)
            label = ("Compliant" if result['is_compliant']
                    else "Lawful requirement not met" if result['unlawful_deficiencies']
                    else "Best practice gap")
            story.append(Paragraph(f"Provision {result['entry_number']} — {label}", h_))
            story.append(Paragraph(f'<i>"{result["entry_text"][:400]}"</i>', body))
            if result['unlawful_deficiencies']:
                story.append(Paragraph("Lawful requirements not met", h3))
                for d in result['unlawful_deficiencies']:
                    story.append(Paragraph(f"⚠ {d}", red_s))
            if result['additional_patterns']:
                story.append(Paragraph("Additional pattern identified", h3))
                for pt in result['additional_patterns']:
                    story.append(Paragraph(f"◈ {pt}", pur_s))
            if result['best_practice_gaps']:
                story.append(Paragraph("Best practice gaps", h3))
                for g in result['best_practice_gaps']:
                    story.append(Paragraph(f"◉ {g}", amb_s))
            if result['ofsted_principle']:
                op = result['ofsted_principle']
                story.append(Paragraph("Inspection framework note", h3))
                story.append(Paragraph(f"<b>{op['area']}:</b> {op['principle']}", amb_s))
            if result['policy_gaps']:
                story.append(Paragraph("School policy cross-reference", h3))
                for pg in result['policy_gaps']:
                    story.append(Paragraph(f"◈ {pg}", pur_s))
            if result['required_specification']:
                story.append(Paragraph("Required specification", h3))
                for spec in result['required_specification']:
                    story.append(Paragraph(f"• {spec}", body))
            if result['tactical_advice']:
                story.append(Paragraph("Tactical advice", h3))
                for advice in result['tactical_advice']:
                    story.append(Paragraph(f"→ {advice}", tac_s))
            if result['unlawful_deficiencies']:
                story.append(Paragraph("If it is not specified and evidenced, it is not lawfully enforceable under the Children and Families Act 2014. Lack of evidence is evidence of lack.", anc_s))
            story.append(Spacer(1, 5*mm))

    story.append(Spacer(1, 6*mm))
    story.append(Paragraph("Upload the expert reports (EP, OT, or SLT) to begin the Cross-Reference report.", tac_s))
    doc.build(story)
    buf.seek(0)
    return buf

# ─────────────────────────────────────────────
# SURVEY
# ─────────────────────────────────────────────

def render_survey():
    st.markdown("---")
    st.markdown("### Beta feedback")
    st.markdown("Takes about two minutes. Every answer goes directly to the team building FRED.")

    with st.form("feedback_form"):
        st.selectbox("Did the report identify anything you did not already know?",
            ["Yes — significantly", "Yes — partially", "No — I knew this already"])
        st.selectbox("Did the traffic light system make sense?",
            ["Yes — very clear", "Mostly clear", "Confusing", "Not sure"])
        st.selectbox("Does the layout feel simple and easy to follow?",
            ["Yes — very simple", "Mostly", "Could be simpler", "No"])
        st.selectbox("How does it look to you?",
            ["Clean and professional", "Fine but nothing special", "Needs more personality", "Not sure"])
        st.selectbox(
            "Would you find it useful to personalise how FRED looks — for example choosing a colour theme or text size?",
            ["Yes — colour theme", "Yes — text size", "Yes — both", "Not bothered", "No"])
        st.selectbox("Would you pay for the one-off report?",
            ["Yes — definitely", "Possibly", "Not sure", "No"])
        st.text_input("What feels like a fair price for the full report?",
            placeholder="e.g. £25, £35, £50...")
        st.selectbox(
            "Would you use a subscription that holds your documents, drafts emails, and prepares you for meetings?",
            ["Yes — definitely", "Possibly", "Not sure", "No"])
        st.text_input("What would feel like a fair monthly price?",
            placeholder="e.g. £10, £15, £20 per month...")
        st.text_area("Anything else — what worked, what did not, what is missing?", height=80)

        st.markdown("---")
        st.markdown("**Would you like to be notified when FRED launches?**")
        notify = st.radio("", ["Yes — notify me", "No thank you"],
                         horizontal=True, label_visibility="collapsed")

        submitted = st.form_submit_button("Submit feedback")
        if submitted:
            st.success(
                "Thank you. Your feedback has been received. "
                "It directly informs the next version of FRED."
            )

# ─────────────────────────────────────────────
# LANDING PAGE
# ─────────────────────────────────────────────

def render_landing():
    st.markdown(f"""
    <div class="fred-nav">
        <div class="fred-nav-logo">FRED</div>
        <div class="fred-nav-links">
            <span class="fred-nav-link">How it works</span>
            <span class="fred-nav-link">Pricing</span>
            <span class="fred-nav-link">About</span>
        </div>
    </div>
    <div class="fred-beta">Beta — design and functionality are actively being developed. Your feedback shapes the final product.</div>
    <div class="fred-hero">
        <div class="fred-big">FRED</div>
        <div class="fred-full">Families' Rights and Entitlements Directory</div>
        <div class="fred-hero-title">Your child's plan should work for <span>your child.</span></div>
        <div class="fred-hero-origin">Built by a parent who learned the hard way — so you don't have to.</div>
        <div class="fred-hero-sub">FRED reads your child's EHCP, identifies every provision that isn't lawfully enforceable, and tells you exactly what to do about it — in plain language, at any hour.</div>
        <div class="fred-hero-flex">Our service is flexible — either a <span>one-off report on your EHCP and provision</span> or a <span>full subscription</span> that holds your child's complete journey.</div>
    </div>
    """, unsafe_allow_html=True)

    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        if st.button("📋  Get my report", key="hero_get_report"):
            st.session_state.stage = 'upload'
            st.rerun()

    st.markdown(f"""
    <div class="fred-hero-cta">
        <div class="fred-btn-reassure">Upload first. Decide after. Your report is ready before you pay.</div>
        <div class="fred-btn-pricing">From £XX for the full report — or <span onclick="">see our subscription plans</span></div>
    </div>
    """, unsafe_allow_html=True)

    st.markdown('<hr class="fred-divider">', unsafe_allow_html=True)

    st.markdown("""
    <div class="fred-section">
        <div class="fred-sec-label">How it works</div>
        <div class="fred-sec-title">Everything you need to know.</div>
        <ul class="fred-bullets">
            <li><span class="fred-bdot"></span>Gather your EHCP, EP report, or any school document — PDF or Word is fine</li>
            <li><span class="fred-bdot"></span>Answer five short questions about your situation — draft or final plan, upcoming dates, how things stand with the school</li>
            <li><span class="fred-bdot"></span>Receive a full report — every provision assessed against the Children and Families Act 2014, the SEND Code of Practice 2015, and your school's own policy where provided</li>
            <li><span class="fred-bdot"></span>Download your report as Word or PDF, whichever works for you</li>
            <li><span class="fred-bdot"></span><span>Then with a subscription you can add school emails, meeting transcripts, and specialist reports to <span class="fred-sub-bullet">build the complete picture</span></span></li>
        </ul>
    </div>
    """, unsafe_allow_html=True)

    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        if st.button("📋  Get my report", key="how_get_report"):
            st.session_state.stage = 'upload'
            st.rerun()
    st.markdown('<div class="fred-upload-wrap"><span class="fred-upload-note">PDF or Word · processed privately · not stored or shared</span></div>', unsafe_allow_html=True)

    st.markdown('<hr class="fred-divider">', unsafe_allow_html=True)

    st.markdown(f"""
    <div class="fred-section">
        <div class="fred-sec-label" style="text-align:center;">The traffic light system</div>
        <div class="fred-sec-title">You'll always know where you stand.</div>
        <div class="fred-sec-sub">Every finding is colour coded so you can see at a glance what matters most.</div>
        <div class="fred-traffic-legend">
            <div class="fred-trow">
                <div class="fred-tdot tdot-red"></div>
                <div class="fred-ttext"><strong>Red — lawful requirement not met.</strong> The provision does not meet the statutory standard set by the Children and Families Act 2014. Must be addressed at annual review.</div>
            </div>
            <div class="fred-trow">
                <div class="fred-tdot tdot-amber"></div>
                <div class="fred-ttext"><strong>Amber — best practice gap.</strong> Meets the minimum lawful standard but falls short of what good practice recommends. Worth raising at annual review.</div>
            </div>
            <div class="fred-trow">
                <div class="fred-tdot tdot-green"></div>
                <div class="fred-ttext"><strong>Green — compliant.</strong> Meets the lawful standard. Use compliant entries as the benchmark when challenging non-compliant ones.</div>
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)

    st.markdown('<hr class="fred-divider">', unsafe_allow_html=True)

    st.markdown("""
    <div class="fred-section">
        <div class="fred-sec-label" style="text-align:center;">Pricing</div>
        <div class="fred-sec-title" style="text-align:center;">Start with what you need.</div>
        <div class="fred-sec-sub">No hidden charges. Your report is ready before you purchase. Every route includes the full report.</div>
        <div class="fred-pricing-grid">
            <div class="fred-price-card">
                <div class="fred-price-name">One-off report</div>
                <div class="fred-price-amount">£XX</div>
                <div class="fred-price-period">single purchase · no commitment</div>
                <ul class="fred-price-features">
                    <li>Full Section F enforceability report</li>
                    <li>Section E SMART outcomes check</li>
                    <li>Traffic light findings — red, amber, green</li>
                    <li>Tactical advice for every finding</li>
                    <li>Downloadable — Word and PDF</li>
                </ul>
            </div>
            <div class="fred-price-card featured">
                <div class="fred-price-badge">Best value</div>
                <div class="fred-price-name">Annual subscription</div>
                <div class="fred-price-amount">£XX</div>
                <div class="fred-price-period">per year · less than £X per week</div>
                <div class="fred-price-first">Includes your report from day one. Year two at a reduced renewal rate.</div>
                <ul class="fred-price-features">
                    <li>Full report included</li>
                    <li>Document vault — all documents held</li>
                    <li>Email support — drafted and calibrated</li>
                    <li>Meeting preparation and script</li>
                    <li>Post-meeting summary emails</li>
                    <li>Annual review preparation pack</li>
                    <li>School transition support</li>
                </ul>
            </div>
            <div class="fred-price-card">
                <div class="fred-price-name">Monthly</div>
                <div class="fred-price-amount">£XX</div>
                <div class="fred-price-period">per month from month two · cancel anytime</div>
                <div class="fred-price-first">First month £XX — includes your report.</div>
                <ul class="fred-price-features">
                    <li>Report included in first month</li>
                    <li>Everything in the full service</li>
                    <li>No annual commitment</li>
                    <li>Cancel anytime</li>
                </ul>
            </div>
        </div>
        <div style="font-size:11px;color:#717D7E;text-align:center;font-style:italic;margin-top:8px;">Prices shown as placeholders — confirmed at launch.</div>
    </div>
    """, unsafe_allow_html=True)

    st.markdown('<hr class="fred-divider">', unsafe_allow_html=True)

    st.markdown("""
    <div class="fred-section">
        <div class="fred-sec-label">From a parent</div>
        <div class="fred-sec-title">You already know something isn't right.</div>
        <div class="fred-quote-box">
            <div class="fred-quote">"I spent three years learning what I should have been told on day one. The language in my son's plan looked like provision. It wasn't. FRED would have shown me that in minutes."</div>
            <div class="fred-quote-attr">— Founder, FRED</div>
        </div>
    </div>
    """, unsafe_allow_html=True)

    st.markdown('<hr class="fred-divider">', unsafe_allow_html=True)

    st.markdown("""
    <div class="fred-section">
        <div class="fred-sec-label">Questions</div>
        <div class="fred-sec-title">Straightforward answers.</div>
        <div class="fred-faq-item">
            <div class="fred-faq-q">Is FRED legal advice?</div>
            <div class="fred-faq-a">No. FRED provides information to help you understand the language of your child's plan and what the law says about it. It does not replace a solicitor or independent advocate. All guidance is referenced to the Children and Families Act 2014 and the SEND Code of Practice 2015. Where you upload a school policy or accessibility plan, FRED cross-references the school's own commitments against the provision in the plan.</div>
        </div>
        <div class="fred-faq-item">
            <div class="fred-faq-q">When do I pay?</div>
            <div class="fred-faq-a">After FRED has read your plan and you have seen a preview of what it found. You upload first, see a finding from your plan, then decide whether to purchase the full report. Nothing is charged until you choose to proceed.</div>
        </div>
        <div class="fred-faq-item">
            <div class="fred-faq-q">Is my data private?</div>
            <div class="fred-faq-a">Yes. For the one-off report your document is read during your session only — not stored or retained. In the full service your documents are held in your own private vault, accessible only to you.</div>
        </div>
        <div class="fred-faq-item">
            <div class="fred-faq-q">What if I don't have an EHCP yet?</div>
            <div class="fred-faq-a">FRED works with EP reports, OT reports, SALT reports, and school correspondence. If you are at the needs assessment stage or have had an assessment refused, FRED can show you what the Children and Families Act 2014 says about your situation and what questions to ask.</div>
        </div>
        <div class="fred-faq-item">
            <div class="fred-faq-q">What documents can I upload?</div>
            <div class="fred-faq-a">Any PDF or Word document — EHCP, EP report, OT report, school emails saved as PDF, meeting transcripts, school SEND policy, behaviour policy, or accessibility plan.</div>
        </div>
        <div class="fred-faq-item">
            <div class="fred-faq-q">Can I cancel my subscription?</div>
            <div class="fred-faq-a">Yes. Monthly subscriptions cancel anytime. Resubscribing resets to the first month rate which includes a fresh report. Annual subscriptions run twelve months. Year two renewals are at a reduced rate as the report is not repeated.</div>
        </div>
    </div>
    """, unsafe_allow_html=True)

    st.markdown('<hr class="fred-divider">', unsafe_allow_html=True)

    st.markdown("""
    <div style="text-align:center;padding:28px 0;background:var(--color-background-secondary);border-radius:10px;margin:16px 0;">
        <div style="font-size:17px;font-weight:500;color:var(--color-text-primary);margin-bottom:6px;">Ready to see what your child's plan actually says?</div>
        <div style="font-size:13px;color:var(--color-text-secondary);margin-bottom:16px;">Upload your document. See a finding. Decide after.</div>
    </div>
    """, unsafe_allow_html=True)

    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        if st.button("📋  Get my report", key="bottom_get_report"):
            st.session_state.stage = 'upload'
            st.rerun()

    st.markdown(f"""
    <div class="fred-footer">
        <div class="fred-footer-logo">FRED</div>
        <div class="fred-footer-text">
            Families' Rights and Entitlements Directory<br>
            Not legal advice · All data private and secure · Built for families navigating the EHCP process<br><br>
            Privacy · Terms · Contact
        </div>
    </div>
    """, unsafe_allow_html=True)

# ─────────────────────────────────────────────
# APP FLOW
# ─────────────────────────────────────────────

if st.session_state.stage == 'landing':
    render_landing()

elif st.session_state.stage == 'upload':

    st.markdown(f"""
    <div class="fred-header-bar">
        <div class="fred-header-title">FRED</div>
        <div class="fred-header-sub">Families' Rights and Entitlements Directory</div>
    </div>
    <div class="fred-beta-notice">
        <strong>Beta v0.5</strong> — Design and functionality are actively being developed.
        Your feedback shapes the final product. FRED provides information to help you
        understand the language of your child's plan and what the law says about it.
        It does not constitute legal advice.
    </div>
    """, unsafe_allow_html=True)

    render_traffic_legend()

    st.markdown("### Get my report")
    st.markdown("Upload any document — EHCP, EP report, school email, meeting transcript, or school policy. PDF or Word. FRED works out what it is.")

    main_file = st.file_uploader(
        "Upload your document",
        type=['pdf', 'docx', 'doc'],
        key='main_upload',
        help="Processed privately during this session only. Not stored or shared."
    )

    st.markdown("""
    <div class="fred-upload-tip">
        <strong>Email:</strong> Open it, select print, choose Save as PDF.<br>
        <strong>Password protected document:</strong> Open it, select print, save as PDF — this removes the lock on most LA documents.
    </div>
    """, unsafe_allow_html=True)

    with st.expander("Add another document — school policy, specialist report, email, or transcript (optional)"):
        st.markdown("FRED cross-references everything you provide. Upload any additional documents here.")
        extra_1 = st.file_uploader("Additional document", type=['pdf','docx','doc'], key='extra_1')
        extra_2 = st.file_uploader("Another document", type=['pdf','docx','doc'], key='extra_2')
        extra_3 = st.file_uploader("Another document", type=['pdf','docx','doc'], key='extra_3')

    if main_file:
        with st.spinner("Fred is reading your document..."):
            text, error = read_file(main_file)
            if error:
                st.error(error)
            else:
                doc_type = detect_doc_type(text)
                sections = identify_sections(text)
                st.session_state.extracted_sections = sections
                st.session_state.raw_text = text

                if sections:
                    st.success(
                        f"Document read. "
                        f"Sections identified: {', '.join(sorted(sections.keys()))}. "
                        f"Key information saved to your report summary."
                    )
                else:
                    st.warning("FRED could not identify standard EHCP sections automatically. Paste Section F text below to proceed.")
                    manual_f = st.text_area("Paste Section F provision text here:", height=180)
                    if manual_f:
                        st.session_state.extracted_sections['F'] = manual_f

        for i, extra in enumerate([extra_1, extra_2, extra_3], 1):
            if extra:
                with st.spinner(f"Reading additional document {i}..."):
                    extra_text, extra_error = read_file(extra)
                    if extra_error:
                        st.warning(f"Document {i}: {extra_error}")
                    else:
                        dtype = detect_doc_type(extra_text)
                        if dtype == 'policy':
                            st.session_state.policy_text = extra_text
                            st.success(f"School policy read — ready for cross-reference.")
                        elif dtype == 'email':
                            st.session_state.email_text = extra_text
                            st.success(f"Email read — ready for correspondence analysis.")
                        elif dtype == 'transcript':
                            st.session_state.transcript_text = extra_text
                            st.success(f"Transcript read — primary record of what was said.")
                        else:
                            st.session_state.policy_text = st.session_state.get('policy_text', '') + ' ' + extra_text
                            st.success(f"Additional document {i} read and added.")

        col1, col2 = st.columns(2)
        with col1:
            if st.button("← Back to home", key="upload_back"):
                st.session_state.stage = 'landing'
                st.rerun()
        with col2:
            if st.button("Continue →", key="upload_continue"):
                st.session_state.stage = 'questions'
                st.rerun()

elif st.session_state.stage == 'questions':

    st.markdown(f"""
    <div class="fred-header-bar">
        <div class="fred-header-title">FRED</div>
        <div class="fred-header-sub">Families' Rights and Entitlements Directory</div>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("### A few quick questions")
    st.markdown("These shape the report you receive.")

    q1 = st.selectbox("1. What have you uploaded?", options=[
        "My child's EHCP",
        "An EP (Educational Psychologist) report",
        "A specialist report (OT, SALT, or other)",
        "School or LA correspondence",
        "Meeting notes or transcript",
        "More than one of the above",
    ])
    st.session_state.answers['q1'] = q1

    if "EHCP" in q1:
        q2 = st.selectbox("2. Is this a draft or final issued EHCP?", options=[
            "Draft — I am still in the review process",
            "Final — this has been formally issued by the LA",
            "I am not sure",
        ])
        st.session_state.answers['q2'] = q2
    else:
        st.session_state.answers['q2'] = "Not an EHCP"

    q3 = st.selectbox("3. Which best describes your situation right now?", options=[
        "I have just received this and want to understand it",
        "I have an upcoming annual review or meeting",
        "I am having difficulty getting the school to deliver what is in the plan",
        "I have had a needs assessment refused",
        "I am just starting the EHCP process",
    ])
    st.session_state.answers['q3'] = q3

    q4 = st.selectbox("4. Do you have any important dates coming up?", options=[
        "No upcoming dates right now",
        "Yes — annual review",
        "Yes — meeting with school or LA",
        "Yes — LA deadline",
    ])
    st.session_state.answers['q4'] = q4
    if q4 != "No upcoming dates right now":
        upcoming_date = st.date_input("When is this?")
        st.session_state.answers['upcoming_date'] = str(upcoming_date)

    q5 = st.selectbox(
        "5. How would you describe your current relationship with the school or LA?",
        options=[
            "Warm and collaborative",
            "Constructive but cautious",
            "Professionally firm",
            "Formally assertive",
            "Rights-based and formal",
        ])
    st.session_state.answers['q5'] = q5

    col1, col2 = st.columns(2)
    with col1:
        if st.button("← Back", key="q_back"):
            st.session_state.stage = 'upload'
            st.rerun()
    with col2:
        if st.button("Run report →", key="q_continue"):
            st.session_state.stage = 'processing'
            st.rerun()

elif st.session_state.stage == 'processing':

    st.markdown(f"""
    <div class="fred-header-bar">
        <div class="fred-header-title">FRED</div>
        <div class="fred-header-sub">Families' Rights and Entitlements Directory</div>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("### Fred is working...")

    sections = st.session_state.extracted_sections
    policy_text = st.session_state.get('policy_text', '')
    email_text = st.session_state.get('email_text', '')
    transcript_text = st.session_state.get('transcript_text', '')

    report_results = []
    section_e_results = []
    correspondence_analysis = None
    post_meeting_email = ''

    with st.spinner("Reading Section F provision entries..."):
        if 'F' in sections:
            entries = extract_entries(sections['F'])
            for i, entry in enumerate(entries):
                if len(entry.strip()) > 20:
                    report_results.append(audit_entry(entry, i+1, policy_text))

    with st.spinner("Checking Section E outcomes..."):
        if 'E' in sections:
            section_e_results = audit_section_e(sections['E'])

    if email_text:
        with st.spinner("Analysing correspondence..."):
            correspondence_analysis = analyse_correspondence(email_text, sections, transcript_text)
            post_meeting_email = generate_post_meeting_email(correspondence_analysis, st.session_state.answers)

    sneak_peek = None
    for r in report_results:
        if r['unlawful_deficiencies']:
            sneak_peek = r
            break
    if sneak_peek is None and report_results:
        sneak_peek = report_results[0]

    st.session_state.report_results = report_results
    st.session_state.section_e_results = section_e_results
    st.session_state.correspondence_analysis = correspondence_analysis
    st.session_state.post_meeting_email = post_meeting_email
    st.session_state.sneak_peek_result = sneak_peek
    st.session_state.stage = 'preview'
    st.rerun()

elif st.session_state.stage == 'preview':

    st.markdown(f"""
    <div class="fred-header-bar">
        <div class="fred-header-title">FRED</div>
        <div class="fred-header-sub">Families' Rights and Entitlements Directory</div>
    </div>
    """, unsafe_allow_html=True)

    sneak_peek = st.session_state.sneak_peek_result

    if sneak_peek:
        render_sneak_peek(sneak_peek)

    st.markdown("---")
    st.markdown("### Get your full report — beta is free")
    st.markdown("FRED is in beta. Enter your email to receive your full report and access the complete service. No payment required during beta.")

    with st.form("email_capture_form"):
        email_input = st.text_input("Your email address", placeholder="your@email.com")
        submitted = st.form_submit_button("Get full access →")
        if submitted:
            if email_input and '@' in email_input:
                st.session_state.email_captured = True
                st.session_state.captured_email = email_input
                st.session_state.stage = 'results'
                st.rerun()
            else:
                st.warning("Please enter a valid email address.")

    st.markdown(f"""
    <div style="font-size:11px;color:{GREY};text-align:center;font-style:italic;margin-top:6px;">
        Your email is used to send your report only. Not shared with third parties.
    </div>
    """, unsafe_allow_html=True)

    col1, col2 = st.columns(2)
    with col1:
        if st.button("← Back", key="preview_back"):
            st.session_state.stage = 'questions'
            st.rerun()

elif st.session_state.stage == 'results':

    st.markdown(f"""
    <div class="fred-header-bar">
        <div class="fred-header-title">FRED</div>
        <div class="fred-header-sub">Families' Rights and Entitlements Directory</div>
    </div>
    """, unsafe_allow_html=True)

    report_results = st.session_state.report_results
    section_e_results = st.session_state.get('section_e_results', [])
    answers = st.session_state.answers
    correspondence_analysis = st.session_state.get('correspondence_analysis')
    post_meeting_email = st.session_state.get('post_meeting_email', '')

    if correspondence_analysis:
        render_correspondence(correspondence_analysis, post_meeting_email)
        st.markdown("---")

    if report_results or section_e_results:
        render_full_report(report_results, section_e_results, answers)

    st.markdown("---")
    st.markdown("### Download your report")
    c1, c2 = st.columns(2)
    with c1:
        docx_buf = generate_docx(report_results, section_e_results, answers)
        st.download_button(
            "⬇ Download as Word (.docx)",
            data=docx_buf,
            file_name="FRED_Report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            help="Best for Windows and Microsoft Office users"
        )
    with c2:
        pdf_buf = generate_pdf(report_results, section_e_results, answers)
        st.download_button(
            "⬇ Download as PDF",
            data=pdf_buf,
            file_name="FRED_Report.pdf",
            mime="application/pdf",
            help="Best for Apple devices — universally readable"
        )

    st.markdown("---")
    st.markdown("### Add more documents")
    st.markdown(
        "Upload school emails, meeting transcripts, specialist reports, or school policies "
        "to build the complete picture. FRED will cross-reference everything."
    )
    extra_post = st.file_uploader(
        "Add a document", type=['pdf', 'docx', 'doc'], key='post_report_upload'
    )
    if extra_post:
        with st.spinner("Reading..."):
            extra_text, extra_error = read_file(extra_post)
            if extra_error:
                st.warning(extra_error)
            else:
                dtype = detect_doc_type(extra_text)
                if dtype == 'email':
                    st.session_state.email_text = extra_text
                    st.success("Email read. Running correspondence analysis...")
                    ca = analyse_correspondence(extra_text, st.session_state.extracted_sections, st.session_state.get('transcript_text', ''))
                    pme = generate_post_meeting_email(ca, answers)
                    st.session_state.correspondence_analysis = ca
                    st.session_state.post_meeting_email = pme
                    st.rerun()
                elif dtype == 'transcript':
                    st.session_state.transcript_text = extra_text
                    st.success("Transcript read and added to the vault.")
                elif dtype == 'policy':
                    st.session_state.policy_text = extra_text
                    st.success("School policy read and added. Re-running report with cross-reference...")
                    sections = st.session_state.extracted_sections
                    if 'F' in sections:
                        entries = extract_entries(sections['F'])
                        new_results = [audit_entry(e, i+1, extra_text) for i, e in enumerate(entries) if len(e.strip()) > 20]
                        st.session_state.report_results = new_results
                    st.rerun()
                else:
                    st.success("Document read and added to your file.")

    render_survey()

    st.markdown("---")
    if st.button("Start new report"):
        for key in list(defaults.keys()):
            if key in st.session_state:
                del st.session_state[key]
        st.rerun()

# ─────────────────────────────────────────────
# FOOTER — NON LANDING PAGES
# ─────────────────────────────────────────────

if st.session_state.stage not in ('landing',):
    st.markdown(
        f"<div style='text-align:center;color:{GREY};font-size:12px;padding-top:8px;'>"
        "FRED — Families' Rights and Entitlements Directory &nbsp;|&nbsp; "
        "Beta v0.5 &nbsp;|&nbsp; Not legal advice &nbsp;|&nbsp; "
        "Documents read during your session only — not stored or retained"
        "</div>",
        unsafe_allow_html=True
    )
